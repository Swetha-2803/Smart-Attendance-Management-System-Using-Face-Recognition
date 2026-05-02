# myapp/views.py
from django.shortcuts import render
from django.utils import timezone
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from .models import Attendance, ODForm, StudentProfile, StaffProfile, Student, StaffAttendance, StaffLeaveForm, StudentLeaveForm
import pandas as pd
from attendance_recognition import start_camera_thread


# ---------------- Home Page ----------------
def index(request):
    # Automatically start camera when server runs
    start_camera_thread()
    return render(request, 'index.html')


# ---------------- Dashboard Page ----------------
@login_required
def attendance_dashboard(request):
    from datetime import timedelta
    
    if hasattr(request.user, 'student_profile'):
        # Student Dashboard Logic
        student = request.user.student_profile.student
        selected_date = request.GET.get('date', timezone.now().date())
        if isinstance(selected_date, str):
            from datetime import datetime
            try:
                selected_date = datetime.strptime(selected_date, '%Y-%m-%d').date()
            except ValueError:
                selected_date = timezone.now().date()
                
        periods = ['period1','period2','period3','period4','period5','period6','period7']
        
        # Period details for selected date
        try:
            record = Attendance.objects.get(student=student, date=selected_date)
            period_details = {}
            today_present = 0
            for p in periods:
                status = getattr(record, p)
                if status in ['Present', 'OD']:
                    today_present += 1
                time_marked = getattr(record, f"{p}_time")
                period_details[p] = {'status': status, 'time': time_marked}
            today_percent = round((today_present / 7) * 100)
        except Attendance.DoesNotExist:
            period_details = {}
            today_percent = 0

        # Calculate Week & Month Percent
        today = timezone.now().date()
        start_of_week = today - timedelta(days=today.weekday())
        start_of_month = today.replace(day=1)
        
        week_records = Attendance.objects.filter(student=student, date__gte=start_of_week, date__lte=today)
        month_records = Attendance.objects.filter(student=student, date__gte=start_of_month, date__lte=today)
        
        def calc_percent(record_set):
            present = 0
            total = len(record_set) * 7
            for r in record_set:
                for p in periods:
                    if getattr(r, p) in ['Present', 'OD']:
                        present += 1
            return round((present / total) * 100) if total > 0 else 0

        approved_ods = ODForm.objects.filter(student=student, status='Approved')
        approved_leaves = StudentLeaveForm.objects.filter(student=student, status='Approved')

        context = {
            'student': student,
            'selected_date': selected_date,
            'period_details': period_details,
            'today_percent': today_percent,
            'week_percent': calc_percent(week_records),
            'month_percent': calc_percent(month_records),
            'approved_ods': approved_ods,
            'approved_leaves': approved_leaves
        }
        return render(request, 'student_dashboard.html', context)
        
    elif hasattr(request.user, 'staff_profile'):
        # Staff Dashboard Logic
        selected_date = request.GET.get('date', timezone.now().date())
        if isinstance(selected_date, str):
            from datetime import datetime
            try:
                selected_date = datetime.strptime(selected_date, '%Y-%m-%d').date()
            except ValueError:
                selected_date = timezone.now().date()

        records = Attendance.objects.select_related('student').filter(date=selected_date)
        pending_ods = ODForm.objects.filter(status='Pending Staff')
        pending_student_leaves = StudentLeaveForm.objects.filter(status='Pending Staff')
        periods = ['period1','period2','period3','period4','period5','period6','period7']

        # Calculate today institution %
        today_present = 0
        today_total = len(records) * 7
        for r in records:
            for p in periods:
                if getattr(r, p) in ['Present', 'OD']:
                    today_present += 1
        today_percent = round((today_present / today_total) * 100) if today_total > 0 else 0

        today = timezone.now().date()
        start_of_week = today - timedelta(days=today.weekday())
        start_of_month = today.replace(day=1)
        
        week_records = Attendance.objects.filter(date__gte=start_of_week, date__lte=today)
        month_records = Attendance.objects.filter(date__gte=start_of_month, date__lte=today)
        
        def calc_percent(record_set):
            present = 0
            total = len(record_set) * 7
            for r in record_set:
                for p in periods:
                    if getattr(r, p) in ['Present', 'OD']:
                        present += 1
            return round((present / total) * 100) if total > 0 else 0

        context = {
            'records': records,
            'selected_date': selected_date,
            'periods': periods,
            'pending_ods': pending_ods,
            'pending_student_leaves': pending_student_leaves,
            'today_percent': today_percent,
            'week_percent': calc_percent(week_records),
            'month_percent': calc_percent(month_records)
        }
        return render(request, 'staff_dashboard.html', context)
        
    elif request.user.is_superuser:
        # Admin Dashboard Logic
        selected_date = request.GET.get('date', timezone.now().date())
        records = Attendance.objects.select_related('student').filter(date=selected_date)
        
        total_students = StudentProfile.objects.count()
        total_staff = StaffProfile.objects.count()
        pending_ods = ODForm.objects.filter(status='Pending HOD')
        total_ods = ODForm.objects.count()
        pending_student_leaves = StudentLeaveForm.objects.filter(status='Pending HOD')
        pending_staff_leaves = StaffLeaveForm.objects.filter(status='Pending')

        # Calculate today's student presents (at least one period present)
        students_present = 0
        periods = ['period1','period2','period3','period4','period5','period6','period7']
        for r in records:
            if any(getattr(r, p) == 'Present' for p in periods):
                students_present += 1
        
        students_absent = total_students - students_present
        
        # Staff placeholders (since we don't track staff faces yet)
        staff_present = total_staff
        staff_absent = 0

        # Calculate Overall Attendance for Long Absentees List
        all_students = Student.objects.all()
        student_stats = []
        total_institution_present = 0
        total_institution_periods = 0

        for student in all_students:
            student_records = Attendance.objects.filter(student=student)
            total_p = 0
            present_p = 0
            for sr in student_records:
                for p in periods:
                    total_p += 1
                    total_institution_periods += 1
                    status = getattr(sr, p)
                    if status == 'Present' or status == 'OD':
                        present_p += 1
                        total_institution_present += 1
            
            percent = (present_p / total_p * 100) if total_p > 0 else 0
            student_stats.append({
                'student': student,
                'percent': round(percent, 1)
            })

        long_absentees = [s for s in student_stats if s['percent'] < 75]
        
        monthly_percent = round((total_institution_present / total_institution_periods * 100), 1) if total_institution_periods > 0 else 0

        # Build daily_records for the table (with percent attached) and sort ascending by percent
        daily_records = []
        for r in records:
            percent = 0
            for s in student_stats:
                if s['student'] == r.student:
                    percent = s['percent']
                    break
            daily_records.append({
                'record': r,
                'percent': percent
            })
        daily_records.sort(key=lambda x: x['percent'])

        context = {
            'records': records,
            'daily_records': daily_records,
            'selected_date': selected_date,
            'periods': periods,
            'total_students': total_students,
            'students_present': students_present,
            'students_absent': students_absent,
            'total_staff': total_staff,
            'staff_present': staff_present,
            'staff_absent': staff_absent,
            'pending_ods': pending_ods,
            'total_ods': total_ods,
            'pending_student_leaves': pending_student_leaves,
            'pending_staff_leaves': pending_staff_leaves,
            'monthly_percent': monthly_percent,
            'long_absentees': long_absentees,
            'student_stats': student_stats
        }
        return render(request, 'admin_dashboard.html', context)
        
    # Fallback to old dashboard if no profile
    return render(request, 'attendance_dashboard.html')

@login_required
def upload_od(request):
    from django.shortcuts import redirect
    from django.contrib import messages
    import PyPDF2
    if request.method == 'POST':
        date = request.POST.get('date')
        reason = request.POST.get('reason')
        doc = request.FILES.get('document')
        rollno = request.POST.get('rollno') # Optional, from staff dashboard
        
        try:
            # Document validation for students
            if hasattr(request.user, 'student_profile') and doc and doc.name.lower().endswith('.pdf'):
                try:
                    pdf_reader = PyPDF2.PdfReader(doc)
                    text = ""
                    for page in pdf_reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted
                    
                    # Check if text was found and reason matches
                    if text.strip() and reason.lower() not in text.lower():
                        messages.error(request, f"Validation Failed: The reason '{reason}' was not found in the uploaded PDF. Please make sure the document contains the stated reason.")
                        return redirect('dashboard')
                except Exception as e:
                    print(f"PDF extraction error: {e}")
                    # Allow fallback to manual verification if extraction completely fails
            
            # Reset file pointer after reading
            if doc:
                doc.seek(0)
                
            if hasattr(request.user, 'staff_profile') or request.user.is_superuser:
                # Staff uploading OD for a student
                if rollno:
                    student = Student.objects.get(student_rollno=rollno)
                    ODForm.objects.create(student=student, date=date, reason=reason, document=doc, status="Approved")
                    messages.success(request, f"OD Form uploaded and auto-approved for Student {rollno}.")
                else:
                    messages.error(request, "Student Roll No is required.")
            
            elif hasattr(request.user, 'student_profile'):
                # Student uploading for themselves
                ODForm.objects.create(
                    student=request.user.student_profile.student,
                    date=date, reason=reason, document=doc
                )
                messages.success(request, "OD Form submitted successfully! Waiting for staff approval.")
        except Exception as e:
            messages.error(request, f"Error uploading OD form: {str(e)}")
            
    return redirect('dashboard')

@login_required
def download_blank_od_form(request):
    import io
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle
    
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    
    # Header
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(width / 2.0, height - 50, "KOVAI KALAIMAGAL EDUCATIONAL TRUST")
    c.setFont("Helvetica-Bold", 12)
    c.drawCentredString(width / 2.0, height - 70, "ATTENDANCE EXEMPTION SLIP")
    
    # Subheader
    c.setFont("Helvetica", 10)
    c.drawCentredString(width / 2.0, height - 90, "CIET    SOACIET    KKCAS    CIMAT")
    
    # Right side box
    c.rect(width - 100, height - 60, 60, 20)
    c.drawString(width - 95, height - 46, "TL26(00)")
    
    # Basic Details
    c.setFont("Helvetica", 11)
    c.drawString(50, height - 130, "Year & Branch: B.Tech (IT) & 3rd Year")
    c.drawString(width - 250, height - 130, "Date: ______________________")
    
    # Table Section
    data = [["S.NO", "NAME", "ROLL NO", "SIGNATURE"]]
    for i in range(1, 11):
        data.append([str(i), "", "", ""])
        
    t = Table(data, colWidths=[40, 200, 100, 150], rowHeights=[25]*11)
    t.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 1, colors.black),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
    ]))
    t.wrapOn(c, width, height)
    t.drawOn(c, 50, height - 430)
    
    # Details Section
    c.drawString(50, height - 480, "Purpose: ________________________________________________________")
    c.drawString(50, height - 520, "Place of Visit: __________________________________________________")
    c.drawString(50, height - 560, "Date & Time: From ______________________ To ______________________")
    
    c.save()
    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="Blank_OD_Form.pdf"'
    return response

@login_required
def approve_od(request, od_id):
    from django.shortcuts import redirect
    from django.contrib import messages
    if request.method == 'POST' and (hasattr(request.user, 'staff_profile') or request.user.is_superuser):
        od = ODForm.objects.get(id=od_id)
        action = request.POST.get('action')
        
        if request.user.is_superuser:
            if action == 'approve':
                od.status = 'Approved'
                od.save()
                # Update Attendance to OD
                att, _ = Attendance.objects.get_or_create(student=od.student, date=od.date)
                for p in ['period1','period2','period3','period4','period5','period6','period7']:
                    if getattr(att, p) == 'Absent':
                        setattr(att, p, 'OD')
                att.save()
                messages.success(request, "OD Form fully approved.")
            elif action == 'reject':
                od.status = 'Rejected'
                od.save()
        else: # Staff Logic
            if action == 'approve':
                od.status = 'Pending HOD'
                od.save()
                messages.success(request, "OD Form approved by Staff. Sent to HOD for final approval.")
            elif action == 'reject':
                od.status = 'Rejected'
                od.save()
    return redirect('dashboard')

@login_required
def upload_student_leave(request):
    from django.shortcuts import redirect
    from django.contrib import messages
    if request.method == 'POST':
        date = request.POST.get('date')
        reason = request.POST.get('reason')
        doc = request.FILES.get('document')
        
        try:
            if hasattr(request.user, 'student_profile'):
                StudentLeaveForm.objects.create(
                    student=request.user.student_profile.student,
                    date=date, reason=reason, document=doc
                )
                messages.success(request, "Leave Form submitted successfully! Waiting for staff approval.")
        except Exception as e:
            messages.error(request, f"Error uploading leave form: {str(e)}")
            
    return redirect('dashboard')

@login_required
def approve_student_leave(request, leave_id):
    from django.shortcuts import redirect
    from django.contrib import messages
    if request.method == 'POST' and (hasattr(request.user, 'staff_profile') or request.user.is_superuser):
        leave = StudentLeaveForm.objects.get(id=leave_id)
        action = request.POST.get('action')
        
        if request.user.is_superuser:
            if action == 'approve':
                leave.status = 'Approved'
                leave.save()
                # Update Attendance to Leave
                att, _ = Attendance.objects.get_or_create(student=leave.student, date=leave.date)
                for p in ['period1','period2','period3','period4','period5','period6','period7']:
                    if getattr(att, p) == 'Absent':
                        setattr(att, p, 'Leave')
                att.save()
                messages.success(request, "Leave approved successfully!")
            elif action == 'reject':
                leave.status = 'Rejected'
                leave.save()
        else: # Staff Logic
            if action == 'approve':
                leave.status = 'Pending HOD'
                leave.save()
                messages.success(request, "Leave Form approved by Staff. Sent to HOD for final approval.")
            elif action == 'reject':
                leave.status = 'Rejected'
                leave.save()
    return redirect('dashboard')

@login_required
def upload_staff_leave(request):
    from django.shortcuts import redirect
    from django.contrib import messages
    if request.method == 'POST' and hasattr(request.user, 'staff_profile'):
        date = request.POST.get('date')
        reason = request.POST.get('reason')
        doc = request.FILES.get('document')
        
        try:
            StaffLeaveForm.objects.create(
                staff=request.user.staff_profile,
                date=date,
                reason=reason,
                document=doc
            )
            messages.success(request, "Leave Permission requested successfully! Waiting for Admin/HOD approval.")
        except Exception as e:
            messages.error(request, f"Error requesting leave: {str(e)}")
            
    return redirect('dashboard')

@login_required
def approve_staff_leave(request, leave_id):
    from django.shortcuts import redirect
    from django.contrib import messages
    if request.method == 'POST' and request.user.is_superuser:
        leave = StaffLeaveForm.objects.get(id=leave_id)
        action = request.POST.get('action')
        if action == 'approve':
            leave.status = 'Approved'
            leave.save()
            # Optionally update StaffAttendance
            att, _ = StaffAttendance.objects.get_or_create(staff=leave.staff, date=leave.date)
            for p in ['period1','period2','period3','period4','period5','period6','period7']:
                if getattr(att, p) == 'Absent':
                    setattr(att, p, 'OD') # Marked as OD/Leave
            att.save()
            messages.success(request, f"Leave for Staff {leave.staff.staff_id} approved.")
        elif action == 'reject':
            leave.status = 'Rejected'
            leave.save()
            messages.warning(request, f"Leave for Staff {leave.staff.staff_id} rejected.")
    return redirect('dashboard')

@login_required
def enroll_user(request):
    from django.shortcuts import redirect
    from django.contrib import messages
    from django.contrib.auth.models import User
    
    if request.method == 'POST' and (hasattr(request.user, 'staff_profile') or request.user.is_superuser):
        role = request.POST.get('role')
        name = request.POST.get('name')
        roll_id = request.POST.get('roll_id')
        image = request.FILES.get('image')
        department = request.POST.get('department', '')
        
        if role and name and roll_id and image:
            try:
                if role == 'Student':
                    student_record = Student.objects.create(
                        student_name=name,
                        student_rollno=roll_id,
                        student_image=image
                    )
                    if not User.objects.filter(username=roll_id).exists():
                        student_user = User.objects.create_user(username=roll_id, password=roll_id)
                        StudentProfile.objects.create(user=student_user, student=student_record)
                    
                    # Create default attendance for today so they show in the dashboard instantly
                    Attendance.objects.get_or_create(
                        student=student_record,
                        date=timezone.now().date(),
                        defaults={
                            "student_name": name,
                            "student_rollno": roll_id,
                            "period1": "Absent", "period2": "Absent",
                            "period3": "Absent", "period4": "Absent",
                            "period5": "Absent", "period6": "Absent",
                            "period7": "Absent"
                        }
                    )
                    messages.success(request, f"Student {name} enrolled successfully!")
                
                elif role == 'Staff':
                    if not User.objects.filter(username=roll_id).exists():
                        staff_user = User.objects.create_user(username=roll_id, password=roll_id)
                        StaffProfile.objects.create(
                            user=staff_user,
                            staff_id=roll_id,
                            department=department,
                            staff_image=image
                        )
                    messages.success(request, f"Staff {name} enrolled successfully!")
            except Exception as e:
                messages.error(request, f"Error enrolling user: {str(e)}")
        else:
            messages.error(request, "Please fill in all required fields.")
            
    return redirect('dashboard')

@login_required
def bulk_enroll(request):
    from django.shortcuts import redirect
    from django.contrib import messages
    from django.contrib.auth.models import User
    import pandas as pd
    
    if request.method == 'POST' and (hasattr(request.user, 'staff_profile') or request.user.is_superuser):
        file = request.FILES.get('bulk_file')
        if not file:
            messages.error(request, "No file uploaded.")
            return redirect('dashboard')
            
        try:
            if file.name.endswith('.csv'):
                df = pd.read_csv(file)
            elif file.name.endswith('.xlsx') or file.name.endswith('.xls'):
                df = pd.read_excel(file)
            else:
                messages.error(request, "Invalid file format. Please upload a CSV or Excel file.")
                return redirect('dashboard')
                
            # Expecting columns: Role, Name, Roll No, Department
            # Normalize column names
            df.columns = [str(c).strip().lower() for c in df.columns]
            
            # Find the actual column names in the dataframe
            role_col = next((c for c in df.columns if 'role' in c), None)
            name_col = next((c for c in df.columns if 'name' in c), None)
            roll_col = next((c for c in df.columns if 'roll' in c or 'id' in c), None)
            dept_col = next((c for c in df.columns if 'dept' in c or 'department' in c), None)
            
            if not (role_col and name_col and roll_col):
                messages.error(request, "Missing required columns in the file. Need 'Role', 'Name', and 'Roll No'.")
                return redirect('dashboard')
                
            success_count = 0
            for index, row in df.iterrows():
                role = str(row[role_col]).strip().title()
                name = str(row[name_col]).strip()
                roll_id = str(row[roll_col]).strip()
                dept = str(row[dept_col]).strip() if dept_col and pd.notna(row[dept_col]) else ""
                
                if not name or not roll_id or str(name) == 'nan' or str(roll_id) == 'nan':
                    continue
                    
                # Create User
                user, user_created = User.objects.get_or_create(username=roll_id, defaults={'password': roll_id})
                if user_created:
                    user.set_password(roll_id)
                    user.save()
                
                if role == 'Student':
                    student, created = Student.objects.get_or_create(
                        student_rollno=roll_id,
                        defaults={'student_name': name}
                    )
                    if created:
                        StudentProfile.objects.create(user=user, student=student)
                        Attendance.objects.get_or_create(
                            student=student,
                            date=timezone.now().date(),
                            defaults={
                                "student_name": name,
                                "student_rollno": roll_id,
                                "period1": "Absent", "period2": "Absent",
                                "period3": "Absent", "period4": "Absent",
                                "period5": "Absent", "period6": "Absent",
                                "period7": "Absent"
                            }
                        )
                        success_count += 1
                        
                elif role == 'Staff':
                    staff, created = StaffProfile.objects.get_or_create(
                        staff_id=roll_id,
                        defaults={
                            'user': user,
                            'department': dept
                        }
                    )
                    if created:
                        success_count += 1
                        
            messages.success(request, f"Successfully enrolled {success_count} users from file!")
            
        except Exception as e:
            messages.error(request, f"Error processing file: {str(e)}")
            
    return redirect('dashboard')

@login_required
def update_attendance(request):
    from django.shortcuts import redirect
    from django.contrib import messages
    
    if request.method == 'POST' and (hasattr(request.user, 'staff_profile') or request.user.is_superuser):
        role = request.POST.get('role', 'Student')
        rollno = request.POST.get('rollno')
        date = request.POST.get('date', timezone.now().date())
        period = request.POST.get('period')
        status = request.POST.get('status')
        
        if rollno and period and status:
            try:
                if role == 'Student':
                    student = Student.objects.get(student_rollno=rollno)
                    att, created = Attendance.objects.get_or_create(student=student, date=date)
                    setattr(att, period, status)
                    att.save()
                    messages.success(request, f"Updated Student {rollno} attendance for {period} to {status}!")
                elif role == 'Staff':
                    staff_profile = StaffProfile.objects.get(staff_id=rollno)
                    att, created = StaffAttendance.objects.get_or_create(staff=staff_profile, date=date)
                    setattr(att, period, status)
                    att.save()
                    messages.success(request, f"Updated Staff {rollno} attendance for {period} to {status}!")
            except Student.DoesNotExist:
                messages.error(request, f"Student with Roll Number {rollno} not found.")
            except StaffProfile.DoesNotExist:
                messages.error(request, f"Staff with ID {rollno} not found.")
            except Exception as e:
                messages.error(request, f"Error updating attendance: {str(e)}")

                
    return redirect('dashboard')


# ---------------- Export to Excel ----------------
def export_attendance_excel(request):
    # Get all attendance records
    date_str = request.GET.get('date', timezone.now().date().strftime('%Y-%m-%d'))
    records = Attendance.objects.filter(date=date_str).values(
        'student_name', 'student_rollno',
        'period1', 'period2', 'period3',
        'period4', 'period5', 'period6', 'period7',
        'date'
    )

    df = pd.DataFrame(records)

    # Add total present column
    if not df.empty:
        df['Number of Periods Present'] = df[
            ['period1','period2','period3','period4','period5','period6','period7']
        ].apply(lambda row: sum(val == "Present" for val in row), axis=1)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="attendance_{date_str}.xlsx"'

    df.to_excel(response, index=False)
    return response

# ---------------- Export to CSV ----------------
def export_attendance_csv(request):
    import csv
    date_str = request.GET.get('date', timezone.now().date().strftime('%Y-%m-%d'))
    records = Attendance.objects.filter(date=date_str)
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="attendance_{date_str}.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['Student Name', 'Roll No', 'Date', 'P1', 'P2', 'P3', 'P4', 'P5', 'P6', 'P7'])
    
    for r in records:
        writer.writerow([r.student_name, r.student_rollno, r.date, r.period1, r.period2, r.period3, r.period4, r.period5, r.period6, r.period7])
        
    return response

# ---------------- Export to Word ----------------
def export_attendance_word(request):
    import docx
    import io
    date_str = request.GET.get('date', timezone.now().date().strftime('%Y-%m-%d'))
    records = Attendance.objects.filter(date=date_str)
    
    doc = docx.Document()
    doc.add_heading(f'Attendance Report - {date_str}', 0)
    
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Roll No'
    hdr_cells[2].text = 'P1'
    hdr_cells[3].text = 'P2'
    hdr_cells[4].text = 'P3'
    hdr_cells[5].text = 'P4'
    hdr_cells[6].text = 'P5'
    hdr_cells[7].text = 'P6'
    hdr_cells[8].text = 'P7'
    
    for r in records:
        row_cells = table.add_row().cells
        row_cells[0].text = r.student_name
        row_cells[1].text = r.student_rollno
        row_cells[2].text = 'P' if r.period1 == 'Present' else 'A' if r.period1 == 'Absent' else r.period1
        row_cells[3].text = 'P' if r.period2 == 'Present' else 'A' if r.period2 == 'Absent' else r.period2
        row_cells[4].text = 'P' if r.period3 == 'Present' else 'A' if r.period3 == 'Absent' else r.period3
        row_cells[5].text = 'P' if r.period4 == 'Present' else 'A' if r.period4 == 'Absent' else r.period4
        row_cells[6].text = 'P' if r.period5 == 'Present' else 'A' if r.period5 == 'Absent' else r.period5
        row_cells[7].text = 'P' if r.period6 == 'Present' else 'A' if r.period6 == 'Absent' else r.period6
        row_cells[8].text = 'P' if r.period7 == 'Present' else 'A' if r.period7 == 'Absent' else r.period7
        
    # Save to memory
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    
    response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="attendance_{date_str}.docx"'
    return response

# ---------------- Export to PDF ----------------
def export_attendance_pdf(request):
    import io
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    
    date_str = request.GET.get('date', timezone.now().date().strftime('%Y-%m-%d'))
    records = Attendance.objects.filter(date=date_str)
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
    elements = []
    
    styles = getSampleStyleSheet()
    elements.append(Paragraph(f"Attendance Report - {date_str}", styles['Title']))
    
    data = [['Student Name', 'Roll No', 'P1', 'P2', 'P3', 'P4', 'P5', 'P6', 'P7']]
    for r in records:
        data.append([
            r.student_name, r.student_rollno,
            'P' if r.period1 == 'Present' else 'A' if r.period1 == 'Absent' else r.period1,
            'P' if r.period2 == 'Present' else 'A' if r.period2 == 'Absent' else r.period2,
            'P' if r.period3 == 'Present' else 'A' if r.period3 == 'Absent' else r.period3,
            'P' if r.period4 == 'Present' else 'A' if r.period4 == 'Absent' else r.period4,
            'P' if r.period5 == 'Present' else 'A' if r.period5 == 'Absent' else r.period5,
            'P' if r.period6 == 'Present' else 'A' if r.period6 == 'Absent' else r.period6,
            'P' if r.period7 == 'Present' else 'A' if r.period7 == 'Absent' else r.period7,
        ])
        
    t = Table(data)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.grey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 12),
        ('BACKGROUND', (0,1), (-1,-1), colors.beige),
        ('GRID', (0,0), (-1,-1), 1, colors.black),
    ]))
    
    elements.append(t)
    doc.build(elements)
    
    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="attendance_{date_str}.pdf"'
    return response
